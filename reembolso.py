import os
import sys
import json
from datetime import date, timedelta
from tkinter import Tk, Button, messagebox, simpledialog
from docx import Document

# =========================
# CAMINHOS E CONFIGURAÇÕES
# =========================

HORARIO_IDA = "07:40"
HORARIO_VOLTA = "17:15"

def get_app_dir():
    base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
    app_dir = os.path.join(base, "Reembolso")
    os.makedirs(app_dir, exist_ok=True)
    return app_dir

CONFIG_FILE = os.path.join(get_app_dir(), "config.json")

def get_docs_dir():
    pasta = os.path.join(os.path.expanduser("~"), "Documents", "Reembolsos")
    os.makedirs(pasta, exist_ok=True)
    return pasta

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# =========================
# CONFIG (DADOS FIXOS)
# =========================

def carregar_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

def salvar_config(dados):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(dados, f, indent=4, ensure_ascii=False)

def pedir_dados():
    nome = simpledialog.askstring("Dados", "Nome completo:")
    cpf = simpledialog.askstring("Dados", "CPF:")
    instituicao = simpledialog.askstring("Dados", "Instituição de ensino:")

    if not nome or not cpf or not instituicao:
        messagebox.showerror("Erro", "Todos os campos são obrigatórios.")
        return

    salvar_config({
        "nome": nome,
        "cpf": cpf,
        "instituicao": instituicao
    })

    messagebox.showinfo("OK", "Dados salvos com sucesso!")

# =========================
# DATAS
# =========================

def dias_uteis(mes, ano):
    d = date(ano, mes, 1)
    dias = []
    while d.month == mes:
        if d.weekday() < 5:  # segunda a sexta
            dias.append(d)
        d += timedelta(days=1)
    return dias

# =========================
# DOCUMENTO
# =========================

def preencher_tabela(tabela, dias):
    linha = 1  # pula cabeçalho
    for dia in dias:
        if linha >= len(tabela.rows):
            tabela.add_row()

        cells = tabela.rows[linha].cells
        cells[0].text = dia.strftime("%d/%m/%Y")
        cells[1].text = HORARIO_IDA
        cells[2].text = HORARIO_VOLTA
        cells[3].text = ""  # assinatura
        linha += 1

def gerar_documento():
    config = carregar_config()
    if not config:
        messagebox.showerror("Erro", "Cadastre os dados pessoais primeiro.")
        return

    mes = simpledialog.askinteger("Mês", "Digite o mês (1-12):", minvalue=1, maxvalue=12)
    ano = simpledialog.askinteger("Ano", "Digite o ano:")

    if not mes or not ano:
        return

    doc = Document(resource_path("modelo.docx"))

    # TEXTO INICIAL
    p = doc.paragraphs[0]
    p.text = (
        f"Eu, {config['nome']}, CPF {config['cpf']}, declaro, para fins de "
        f"recebimento de auxílio, que frequentei presencialmente as aulas na "
        f"instituição de ensino {config['instituicao']} e utilizei de transporte, "
        f"conforme datas e horários abaixo."
    )

    # REEMBOLSO
    doc.paragraphs[1].text = f"Reembolso referente ao mês: {mes}/{ano}"

    # TABELA EXISTENTE
    tabela = doc.tables[0]

    dias_confirmados = []
    for dia in dias_uteis(mes, ano):
        if messagebox.askyesno(
            "Presença",
            f"Você foi no dia {dia.strftime('%d/%m/%Y')}?"
        ):
            dias_confirmados.append(dia)

    if not dias_confirmados:
        messagebox.showinfo("Aviso", "Nenhum dia selecionado.")
        return

    preencher_tabela(tabela, dias_confirmados)

    caminho = os.path.join(
        get_docs_dir(),
        f"Reembolso_{mes}_{ano}.docx"
    )
    doc.save(caminho)

    messagebox.showinfo(
        "Sucesso",
        f"Arquivo gerado em:\n{caminho}"
    )

# =========================
# INTERFACE
# =========================

root = Tk()
root.title("Gerador de Reembolso")
root.geometry("300x220")
root.resizable(False, False)

Button(root, text="Gerar Reembolso", width=30, command=gerar_documento).pack(pady=20)
Button(root, text="Alterar dados pessoais", width=30, command=pedir_dados).pack()

if not carregar_config():
    pedir_dados()

root.mainloop()
